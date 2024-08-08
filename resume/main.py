# from unittest import main
# from xml.dom.minidom import Document
# import docx
# import datetime
# import Doc
# my_doc = docx.Document("HI MY NAME.docx")
# all_paras = my_doc.paragraphs
# print(len(all_paras))
# for x in (all_paras):
#     print(type(x.text))
#     print(x.text)
# specific_para = my_doc.paragraphs[-2]
# print(specific_para.text)
# print(type(specific_para))
# for run in specific_para.runs:
#     print(run.text)
# my_doc.add_paragraph("IMPACT 205 is high ranked failures")
# specific_para = ""
# # specific_para = my_doc.paragraphs[-1]
# var1 = specific_para.add_run("Harry")
# var1.font.name = "Harry P"
# my_doc.save("HI MY NAME.docx")
# p2 = my_doc.add_paragraph('')
# run2 = p2.add_run('this is Harry Potter')
# run2.font.name = 'Harry P'
# my_doc.save("HI MY NAME.docx")
# all_paras = my_doc.paragraphs
# for x in (all_paras):
#     print(x.text)
# document = Document()
# paragraph = document.add_pragraph('Lorem ipsum dolor sit amet.')
# document.add_heading('the REAL meaning of the universe')
# document.add_heading('The role of dolphins', level=1)
# document.save("HI MY NAME.docx")
# for paragraph in document.paragraphs:
#     if 'dolphins' in paragraph.text:
#         paragraph.text = paragraph.text.replace('dolphins', 'Rabbit')
# document.save('HI MY NAME.docx')
# def change_table_values(param, my_date, param1, param2, param3):
#     pass
# def save_doc(st_name, teacher_name, learning_aims, date_sub, comment, deadline):
#     change_table_values("White And Brown Vintage Resume (1).pdf", st_name, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", teacher_name, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", learning_aims, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", date_sub, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", deadline, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", comment, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", st_name, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", date_sub, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", teacher_name, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", teacher_name, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", my_date, 1, 4, f'{st_name}.docx')
#     change_table_values(f"files/{st_name}.docx", my_date, 1, 4, f'{st_name}.docx')
# import csv
# import datetime
# x = datetime.datetime.now()
# my_date = x.strftime("%d/%B/%y")
# with open("HI MY NAME.docx", "r", newline="") as file:
#  csv_reader = csv.reader(file)
# for x in csv_reader:
#         if x[0] == "id":
#             continue
#         save_doc(f"{x[1]} {x[2]}","Fotima Shavkatova", "L01 L02", x[4], f"Grade {x[3]} \n")
# from io import BytesIO
#
# # 7011868644:AAGEzYTDLAVpcf2NlUHHUPLljwotyt3UJYo
# import telebot
# from docx import Document
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.shared import Inches, Pt, RGBColor
# import requests
# import os
#
# # Initialize your bot with the token
# bot = telebot.TeleBot("7011868644:AAGEzYTDLAVpcf2NlUHHUPLljwotyt3UJYo")
#
# # Dictionary to store user data temporarily
# user_data = {}
#
# # Function to handle the /start command
# @bot.message_handler(commands=['start'])
# def start(message):
#     bot.reply_to(message, "Welcome to CV Generator Bot! Please upload your photo.")
#
# # Function to handle the photo upload
# @bot.message_handler(content_types=['photo'])
# def handle_photo(message):
#     chat_id = message.chat.id
#     bot.send_message(chat_id, "Photo uploaded successfully! Now, please provide the following information:\n"
#                               "1. Name")
#
#     photo = message.photo[-1]  # Get the last photo from the list of photos
#     photo_info = bot.get_file(photo.file_id)
#     photo_url = f"https://api.telegram.org/file/bot{bot.token}/{photo_info.file_path}"
#     saved_photo_path = download_photo(photo_url)
#     if saved_photo_path:
#         user_data[chat_id] = {'Photo': saved_photo_path}
#     else:
#         bot.send_message(chat_id, "Failed to upload photo. Please try again.")
#
# # Function to download the photo from the URL
# def download_photo(url):
#     try:
#         response = requests.get(url)
#         if response.status_code == 200:
#             photo_path = "temp_photo.jpg"
#             with open(photo_path, 'wb') as f:
#                 f.write(response.content)
#             return photo_path
#         else:
#             return None
#     except Exception as e:
#         print("Error downloading photo:", e)
#         return None
#
# # Function to handle user messages and gather CV information
# @bot.message_handler(func=lambda message: True)
# def handle_message(message):
#     chat_id = message.chat.id
#     if chat_id not in user_data:
#         bot.send_message(chat_id, "Please start by uploading your photo.")
#         return
#
#     step = len(user_data[chat_id]) + 1
#     if step == 1:
#         user_data[chat_id]['Name'] = message.text
#         bot.send_message(chat_id, '2. Surname')
#     elif step == 2:
#         user_data[chat_id]['Surname'] = message.text
#         bot.send_message(chat_id, '3. Age')
#     elif step == 3:
#         user_data[chat_id]['Age'] = message.text
#         bot.send_message(chat_id, '4. Address')
#     elif step == 4:
#         user_data[chat_id]['Address'] = message.text
#         bot.send_message(chat_id, '5. Education')
#     elif step == 5:
#         user_data[chat_id]['Education'] = message.text
#         bot.send_message(chat_id, '6. Specialization')
#     elif step == 6:
#         user_data[chat_id]['Specialization'] = message.text
#         bot.send_message(chat_id, '7. Skills')
#     elif step == 7:
#         user_data[chat_id]['Skills'] = message.text
#         bot.send_message(chat_id, '8. Languages')
#     elif step == 8:
#         user_data[chat_id]['Languages'] = message.text
#         bot.send_message(chat_id, '9. Experience')
#     elif step == 9:
#         user_data[chat_id]['Experience'] = message.text
#         create_cv(message)
#
# # Function to create the CV
# # def create_cv(message):
# #     chat_id = message.chat.id
# #     cv_doc = Document()
# #
# #     # Add photo to the document
# #     if 'Photo' in user_data[chat_id]:
# #         photo_path = user_data[chat_id]['Photo']
# #         cv_doc.add_picture(photo_path, width=Inches(2.0))
# #         os.remove(photo_path)  # Remove the temporary photo file
# #
# #     # Add header to the document
# #     cv_doc.add_heading('Curriculum Vitae', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# #
# #     # Add user data to the document
# #     for key, value in user_data[chat_id].items():
# #         if key != 'Photo':
# #             cv_doc.add_paragraph(f"{key}: ", style='Heading 2').add_run(value).bold = True
# #
# #     # Generate a clean file name without special characters
# #     clean_name = "".join([c for c in user_data[chat_id]['Name'] if c.isalnum()])
# #     clean_surname = "".join([c for c in user_data[chat_id]['Surname'] if c.isalnum()])
# #     updated_file_name = f"{clean_name}_{clean_surname}_Updated_CV.docx"
# #
# #     cv_doc.save(updated_file_name)
# #
# #     with open(updated_file_name, 'rb') as doc:
# #         bot.send_document(chat_id, doc)
# #
# #     user_data.pop(chat_id, None)
# def create_cv(message):
#     # Name the base document
#     file_name = "ass9CS.docx"
#
#     try:
#         # Load the base document
#         cv_doc = Document(file_name)
#     except FileNotFoundError:
#         bot.send_message(message.chat.id, "The base document 'Resume.docx' was not found.")
#         return
#
#     # Adding the photo to the document
#     if 'Photo' in user_data[message.chat.id]:
#         photo_stream = BytesIO(user_data[message.chat.id]['Photo'])
#         # Insert the photo at the beginning of the document
#         cv_doc.add_picture(photo_stream)
#
#     # Add other information to the document
#     for key, value in user_data[message.chat.id].items():
#         if key != 'Photo':
#             p = cv_doc.add_paragraph()
#             p.add_run(f"{key}: ").bold = True  # Set text to bold
#             p.add_run(value).italic = True  # Set text to italic
#             # Apply formatting as needed
#             for run in p.runs:
#                 run.font.size = Pt(18)
#                 break
#             p.runs[-1].font.size = Pt(18)  # Font size for the value
#             p.runs[-1].font.color.rgb = RGBColor(0, 0, 255)  # Blue color for the value
#
#     # Save the updated CV document
#     updated_file_name = f"{user_data[message.chat.id]['Name']}_{user_data[message.chat.id]['Surname']}_Updated_CV.docx"
#     cv_doc.save(updated_file_name)
#
#     # Send the generated CV document back to the user
#     with open(updated_file_name, 'rb') as doc:
#         bot.send_document(message.chat.id, doc)
#
#     # Clear user data after processing
#     user_data.pop(message.chat.id, None)
# # Start the bot
# bot.polling()
import logging

import telebot
from telebot.types import InputMediaPhoto
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests
from io import BytesIO
import os

TOKEN = "7011868644:AAGEzYTDLAVpcf2NlUHHUPLljwotyt3UJYo"
bot = telebot.TeleBot(TOKEN)

user_data = {}

class UserData:
    def __init__(self):
        self.name = None
        self.surname = None
        self.age = None
        self.address = None
        self.objective = None
        self.education = None
        self.specialization = None
        self.skills = None
        self.languages = None
        self.photo_id = None
        self.experience = ""

def create_modern_cv(chat_id):
    try:
        user = user_data[chat_id]
        doc = Document()
        setup_document_style(doc)

        header = doc.add_heading(level=0)
        run = header.add_run('RESUME')
        run.bold = True
        run.font.size = Pt(35)
        run.font.color.rgb = RGBColor(255, 55, 25)  # Soft purple
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # Adding extra space

        if user.photo_id:
            add_photo(doc, user.photo_id)

        if user.objective:
            add_objective_section(doc, user.objective)

        add_info_section(doc, ' Information', [
            f" user Name: {user.name} {user.surname}",
            f" user Age: {user.age}",
            f" user Address: {user.address}"
        ], 'FFF1')  # Lavender blush

        add_info_section(doc, 'Education info ', [user.education], 'FFF1')

        add_info_section(doc, 'Skills info', user.skills.split(','), 'FFF1', True)

        add_info_section(doc, 'Languages info', user.languages.split(','), 'FFF1', True)

        add_info_section(doc, 'Professional Experience', [user.experience], 'FFF1')

        doc_path = f'cv_{chat_id}.docx'
        doc.save(doc_path)
        with open(doc_path, 'rb') as doc_file:
            bot.send_document(chat_id, doc_file)
        del user_data[chat_id]  # Deleting user data after sending CV
    except Exception as e:
        bot.send_message(chat_id, f"An error occurred: {str(e)}")

def setup_document_style(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def add_photo(doc, photo_id):
    file_info = bot.get_file(photo_id)
    file_path = f"https://api.telegram.org/file/bot{TOKEN}/{file_info.file_path}"
    image_stream = BytesIO(requests.get(file_path).content)
    doc.add_picture(image_stream, width=Inches(2))
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_objective_section(doc, objective):
    objective_header = doc.add_heading('Objective', level=1)
    objective_para = doc.add_paragraph(objective)
    add_colored_background(objective_para, 'FFF1')  # Lavender blush

def add_info_section(doc, title, content, background_color, bullet_list=True):
    section = doc.add_heading(title, level=1)
    for item in content:
        para = doc.add_paragraph(item, style='ListBullet' if bullet_list else None)
        add_colored_background(para, background_color)

def add_colored_background(paragraph, hex_color):
    shading_el = OxmlElement('w:shd')
    shading_el.set(qn('w:fill'), hex_color)
    paragraph._p.get_or_add_pPr().append(shading_el)

@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    msg = bot.reply_to(message, "Welcome to the CV Bot. Please enter your name.")
    bot.register_next_step_handler(msg, process_name_step)

def process_name_step(message):
    chat_id = message.chat.id
    user_data[chat_id] = UserData()
    user_data[chat_id].name = message.text
    msg = bot.reply_to(message, "1. What is your surname?")
    bot.register_next_step_handler(msg, process_surname_step)

def process_surname_step(message):
    chat_id = message.chat.id
    user_data[chat_id].surname = message.text
    msg = bot.reply_to(message, "2.How old are you?")
    bot.register_next_step_handler(msg, process_age_step)

def process_age_step(message):
    chat_id = message.chat.id
    user_data[chat_id].age = message.text
    msg = bot.reply_to(message, "3.What is your address?")
    bot.register_next_step_handler(msg, process_address_step)

def process_address_step(message):
    chat_id = message.chat.id
    user_data[chat_id].address = message.text
    msg = bot.reply_to(message, "4.What is your objective?")
    bot.register_next_step_handler(msg, process_objective_step)

def process_objective_step(message):
    chat_id = message.chat.id
    user_data[chat_id].objective = message.text
    msg = bot.reply_to(message, "5.about your education.")
    bot.register_next_step_handler(msg, process_education_step)

def process_education_step(message):
    chat_id = message.chat.id
    user_data[chat_id].education = message.text
    msg = bot.reply_to(message, "6. your specialization?")
    bot.register_next_step_handler(msg, process_specialization_step)

def process_specialization_step(message):
    chat_id = message.chat.id
    user_data[chat_id].specialization = message.text
    msg = bot.reply_to(message, "7. How skills do you have? by commas")
    bot.register_next_step_handler(msg, process_skills_step)

def process_skills_step(message):
    chat_id = message.chat.id
    user_data[chat_id].skills = message.text
    msg = bot.reply_to(message, "8. What languages do you speak?  by commas")
    bot.register_next_step_handler(msg, process_languages_step)

def process_languages_step(message):
    chat_id = message.chat.id
    user_data[chat_id].languages = message.text
    msg = bot.reply_to(message, "9. Please  your work experience. (send 'finish' when finished)")
    bot.register_next_step_handler(msg, process_experience_step)

def process_experience_step(message):
    chat_id = message.chat.id
    if message.text.lower() != 'finish':
        user_data[chat_id].experience += "\n" + message.text
        msg = bot.reply_to(message, "Add another item of work experience, or type 'finish' if you are finished.")
        bot.register_next_step_handler(msg, process_experience_step)
    else:
        msg = bot.reply_to(message, " upload your photo.")
        bot.register_next_step_handler(msg, handle_photo)

def handle_photo(message):
    chat_id = message.chat.id
    if 'photo' in message.json:
        photo = message.photo[-1]
        file_id = photo.file_id
        user_data[chat_id].photo_id = file_id
        create_modern_cv(chat_id)
    else:
        msg = bot.reply_to(message, "No photo detected, please try again by uploading a photo.")
        bot.register_next_step_handler(msg, handle_photo)

bot.polling(none_stop=True)





# import telebot
# from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton
# from docx import Document
# from docx.shared import Pt, Inches, RGBColor
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from io import BytesIO
# import requests
#
# TOKEN = "7011868644:AAGEzYTDLAVpcf2NlUHHUPLljwotyt3UJYo"
# bot = telebot.TeleBot(TOKEN)
#
# user_data = {}
#
#
# class UserData:
#     def __init__(self):
#         self.name = None
#         self.surname = None
#         self.age = None
#         self.address = None
#         self.objective = None
#         self.education = None
#         self.specialization = None
#         self.skills = None
#         self.languages = None
#         self.photo_id = None
#         self.experience = ""
#
#
# def create_modern_cv(chat_id):
#     try:
#         user = user_data[chat_id]
#         doc = Document()
#         setup_document_style(doc)
#
#         header = doc.add_heading(level=0)
#         run = header.add_run('RESUME')
#         run.bold = True
#         run.font.size = Pt(35)
#         run.font.color.rgb = RGBColor(255, 55, 25)  # Soft purple
#         header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         doc.add_paragraph()  # Adding extra space
#
#         if user.photo_id:
#             add_photo(doc, user.photo_id)
#
#         if user.objective:
#             add_objective_section(doc, user.objective)
#
#         add_info_section(doc, 'Information', [
#             f"User Name: {user.name} {user.surname}",
#             f"User Age: {user.age}",
#             f"User Address: {user.address}"
#         ], 'FFF1')  # Lavender blush
#
#         add_info_section(doc, 'Education info', [user.education], 'FFF1')
#
#         add_info_section(doc, 'Skills info', user.skills.split(','), 'FFF1', True)
#
#         add_info_section(doc, 'Languages info', user.languages.split(','), 'FFF1', True)
#
#         add_info_section(doc, 'Professional Experience', [user.experience], 'FFF1')
#
#         doc_path = f'cv_{chat_id}.docx'
#         doc.save(doc_path)
#
#         # Send the document
#         with open(doc_path, 'rb') as doc_file:
#             bot.send_document(chat_id, doc_file)
#
#         # Delete user data after sending CV
#         del user_data[chat_id]
#     except Exception as e:
#         bot.send_message(chat_id, f"An error occurred: {str(e)}")
#
#
# def setup_document_style(doc):
#     sections = doc.sections
#     for section in sections:
#         section.top_margin = Inches(1)
#         section.bottom_margin = Inches(1)
#         section.left_margin = Inches(1)
#         section.right_margin = Inches(1)
#
#
# def add_photo(doc, photo_id):
#     file_info = bot.get_file(photo_id)
#     file_path = f"https://api.telegram.org/file/bot{TOKEN}/{file_info.file_path}"
#     image_stream = BytesIO(requests.get(file_path).content)
#     doc.add_picture(image_stream, width=Inches(2))
#     doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#
# def add_objective_section(doc, objective):
#     objective_header = doc.add_heading('Objective', level=1)
#     objective_para = doc.add_paragraph(objective)
#     add_colored_background(objective_para, 'FFF1')  # Lavender blush
#
#
# def add_info_section(doc, title, content, background_color, bullet_list=True):
#     section = doc.add_heading(title, level=1)
#     for item in content:
#         para = doc.add_paragraph(item, style='List Bullet' if bullet_list else None)
#         add_colored_background(para, background_color)
#
#
#
# def add_colored_background(paragraph, hex_color):
#     shading_el = paragraph._element
#     shading_el.get_or_add_shading().val = "clear"
#     shading_el.get_or_add_shading().color = hex_color
#
#
# @bot.message_handler(commands=['start', 'help'])
# def send_welcome(message):
#     msg = bot.reply_to(message, "Welcome to the CV Bot. Please enter your name.")
#     bot.register_next_step_handler(msg, process_name_step)
#
#
# def process_name_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id] = UserData()
#     user_data[chat_id].name = message.text
#     msg = bot.reply_to(message, "1. What is your surname?")
#     bot.register_next_step_handler(msg, process_surname_step)
#
#
# def process_surname_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].surname = message.text
#     msg = bot.reply_to(message, "2. How old are you?")
#     bot.register_next_step_handler(msg, process_age_step)
#
#
# def process_age_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].age = message.text
#     msg = bot.reply_to(message, "3. What is your address?")
#     bot.register_next_step_handler(msg, process_address_step)
#
#
# def process_address_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].address = message.text
#     msg = bot.reply_to(message, "4. What is your objective?")
#     bot.register_next_step_handler(msg, process_objective_step)
#
#
# def process_objective_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].objective = message.text
#     msg = bot.reply_to(message, "5. Tell me about your education.")
#     bot.register_next_step_handler(msg, process_education_step)
#
#
# def process_education_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].education = message.text
#     msg = bot.reply_to(message, "6. What is your specialization?")
#     bot.register_next_step_handler(msg, process_specialization_step)
#
#
# def process_specialization_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].specialization = message.text
#     msg = bot.reply_to(message, "7. How many skills do you have? (separated by commas)")
#     bot.register_next_step_handler(msg, process_skills_step)
#
#
# def process_skills_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].skills = message.text
#     msg = bot.reply_to(message, "8. What languages do you speak? (separated by commas)")
#     bot.register_next_step_handler(msg, process_languages_step)
#
#
# def process_languages_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].languages = message.text
#     msg = bot.reply_to(message, "9. Please list your work experience. (send 'finish' when finished)")
#     bot.register_next_step_handler(msg, process_experience_step)
#
#
# def process_experience_step(message, handle_photo=None):
#     chat_id = message.chat.id
#     if message.text.lower() != 'finish':
#         user_data[chat_id].experience += "\n" + message.text
#         msg = bot.reply_to(message, "Add another item of work experience, or type 'finish' if you are finished.")
#         bot.register_next_step_handler(msg, process_experience_step)
#     else:
#         msg = bot.reply_to(message, "Upload your photo.")
#         bot.register_next_step_handler(msg, handle_photo)
#
#     # Create inline keyboard
#     keyboard = InlineKeyboardMarkup()
#     keyboard.add(InlineKeyboardButton("Personal Info", callback_data='personal_info'))
#     keyboard.add(InlineKeyboardButton("Education", callback_data='education'))
#     keyboard.add(InlineKeyboardButton("Skills", callback_data='skills'))
#     keyboard.add(InlineKeyboardButton("Languages", callback_data='languages'))
#     keyboard.add(InlineKeyboardButton("Experience", callback_data='experience'))
#
#     bot.send_message(chat_id, "Choose a section:", reply_markup=keyboard)
#
#
# @bot.callback_query_handler(func=lambda call: True)
# def callback_query(call):
#     chat_id = call.message.chat.id
#     if call.data == 'personal_info':
#         send_personal_info(chat_id)
#     elif call.data == 'education':
#         send_education_info(chat_id)
#     elif call.data == 'skills':
#         send_skills_info(chat_id)
#     elif call.data == 'languages':
#         send_languages_info(chat_id)
#     elif call.data == 'experience':
#         send_experience_info(chat_id)
#
#
# def send_personal_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id,
#                          f"Personal Info:\nName: {user.name} {user.surname}\nAge: {user.age}\nAddress: {user.address}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
#
# def send_education_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id, f"Education Info:\n{user.education}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
#
# def send_skills_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id, f"Skills Info:\n{', '.join(user.skills.split(','))}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
#
# def send_languages_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id, f"Languages Info:\n{', '.join(user.languages.split(','))}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
#
# def send_experience_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id, f"Experience:\n{user.experience}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
# def create_cv(message):
#     try:
#         cv_doc = Document()
#     except FileNotFoundError:
#         bot.send_message(message.chat.id, "The base document 'Resume.docx' was not found.")
#         return
#
#     chat_id = message.chat.id
#
#     # Add photo to the document
#     if 'Photo' in user_data[chat_id]:
#         photo_emoji = user_data[chat_id]['Photo']
#         cv_doc.add_paragraph(f"Photo: {photo_emoji}")
#
#     # Add header to the document
#     cv_doc.add_heading('Curriculum Vitae', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#     # Add user data to the document
#     for key, value in user_data[chat_id].items():
#         if key != 'Photo':
#             cv_doc.add_paragraph(f"{key}: {value}")
#
#     # Generate a clean file name without special characters
#     clean_name = "".join([c for c in user_data[chat_id]['Name'] if c.isalnum()])
#     clean_surname = "".join([c for c in user_data[chat_id]['Surname'] if c.isalnum()])
#     updated_file_name = f"{clean_name}_{clean_surname}_Updated_CV.docx"
#
#     cv_doc.save(updated_file_name)
#
#     with open(updated_file_name, 'rb') as doc:
#         bot.send_document(chat_id, doc)
#
#     user_data.pop(chat_id, None)
# bot.polling(none_stop=True)
#

#
# import telebot
# from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton
# from docx import Document
# from docx.shared import Pt, Inches, RGBColor
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from io import BytesIO
# import requests
#
# TOKEN = "7011868644:AAGEzYTDLAVpcf2NlUHHUPLljwotyt3UJYo"
# bot = telebot.TeleBot(TOKEN)
#
# user_data = {}
#
#
# class UserData:
#     def __init__(self):
#         self.name = None
#         self.surname = None
#         self.age = None
#         self.address = None
#         self.objective = None
#         self.education = None
#         self.specialization = None
#         self.skills = None
#         self.languages = None
#         self.photo_id = None
#         self.experience = ""
#
#
# def create_modern_cv(chat_id):
#     try:
#         user = user_data[chat_id]
#         doc = Document()
#         setup_document_style(doc)
#
#         header = doc.add_heading(level=0)
#         run = header.add_run('RESUME')
#         run.bold = True
#         run.font.size = Pt(35)
#         run.font.color.rgb = RGBColor(255, 55, 25)  # Soft purple
#         header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         doc.add_paragraph()  # Adding extra space
#
#         if user.photo_id:
#             add_photo(doc, user.photo_id)
#
#         if user.objective:
#             add_objective_section(doc, user.objective)
#
#         add_info_section(doc, 'Information', [
#             f"User Name: {user.name} {user.surname}",
#             f"User Age: {user.age}",
#             f"User Address: {user.address}"
#         ], 'FFF1')  # Lavender blush
#
#         add_info_section(doc, 'Education info', [user.education], 'FFF1')
#
#         add_info_section(doc, 'Skills info', user.skills.split(','), 'FFF1', True)
#
#         add_info_section(doc, 'Languages info', user.languages.split(','), 'FFF1', True)
#
#         add_info_section(doc, 'Professional Experience', [user.experience], 'FFF1')
#
#         doc_path = f'cv_{chat_id}.docx'
#         doc.save(doc_path)
#
#         # Send the document
#         with open(doc_path, 'rb') as doc_file:
#             bot.send_document(chat_id, doc_file)
#
#         # Delete user data after sending CV
#         del user_data[chat_id]
#     except Exception as e:
#         bot.send_message(chat_id, f"An error occurred: {str(e)}")
#
#
# def setup_document_style(doc):
#     sections = doc.sections
#     for section in sections:
#         section.top_margin = Inches(1)
#         section.bottom_margin = Inches(1)
#         section.left_margin = Inches(1)
#         section.right_margin = Inches(1)
#
#
# def add_photo(doc, photo_id):
#     file_info = bot.get_file(photo_id)
#     file_path = f"https://api.telegram.org/file/bot{TOKEN}/{file_info.file_path}"
#     image_stream = BytesIO(requests.get(file_path).content)
#     doc.add_picture(image_stream, width=Inches(2))
#     doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#
# def add_objective_section(doc, objective):
#     objective_header = doc.add_heading('Objective', level=1)
#     objective_para = doc.add_paragraph(objective)
#     add_colored_background(objective_para, 'FFF1')  # Lavender blush
#
#
# def add_info_section(doc, title, content, background_color, bullet_list=True):
#     section = doc.add_heading(title, level=1)
#     for item in content:
#         para = doc.add_paragraph(item, style='List Bullet' if bullet_list else None)
#         add_colored_background(para, background_color)
#
#
# def add_colored_background(paragraph, hex_color):
#     shading_el = paragraph._element
#     shading_el.get_or_add_shading().val = "clear"
#     shading_el.get_or_add_shading().color = hex_color
#
#
# @bot.message_handler(commands=['start', 'help'])
# def send_welcome(message):
#     msg = bot.reply_to(message, "Welcome to the CV Bot. Please enter your name.")
#     bot.register_next_step_handler(msg, process_name_step)
#
#
# def process_name_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id] = UserData()
#     user_data[chat_id].name = message.text
#     msg = bot.reply_to(message, "1. What is your surname?")
#     bot.register_next_step_handler(msg, process_surname_step)
#
#
# def process_surname_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].surname = message.text
#     msg = bot.reply_to(message, "2. How old are you?")
#     bot.register_next_step_handler(msg, process_age_step)
#
#
# def process_age_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].age = message.text
#     msg = bot.reply_to(message, "3. What is your address?")
#     bot.register_next_step_handler(msg, process_address_step)
#
#
# def process_address_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].address = message.text
#     msg = bot.reply_to(message, "4. What is your objective?")
#     bot.register_next_step_handler(msg, process_objective_step)
#
#
# def process_objective_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].objective = message.text
#     msg = bot.reply_to(message, "5. Tell me about your education.")
#     bot.register_next_step_handler(msg, process_education_step)
#
#
# def process_education_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].education = message.text
#     msg = bot.reply_to(message, "6. What is your specialization?")
#     bot.register_next_step_handler(msg, process_specialization_step)
#
#
# def process_specialization_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].specialization = message.text
#     msg = bot.reply_to(message, "7. How many skills do you have? (separated by commas)")
#     bot.register_next_step_handler(msg, process_skills_step)
#
#
# def process_skills_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].skills = message.text
#     msg = bot.reply_to(message, "8. What languages do you speak? (separated by commas)")
#     bot.register_next_step_handler(msg, process_languages_step)
#
#
# def process_languages_step(message):
#     chat_id = message.chat.id
#     user_data[chat_id].languages = message.text
#     msg = bot.reply_to(message, "9. Please list your work experience. (send 'finish' when finished)")
#     bot.register_next_step_handler(msg, process_experience_step)
#
#
# def process_experience_step(message, handle_photo=None):
#     chat_id = message.chat.id
#     if message.text.lower() != 'finish':
#         user_data[chat_id].experience += "\n" + message.text
#         msg = bot.reply_to(message, "Add another item of work experience, or type 'finish' if you are finished.")
#         bot.register_next_step_handler(msg, process_experience_step)
#     else:
#         msg = bot.reply_to(message, "Upload your photo.")
#         bot.register_next_step_handler(msg, handle_photo)
#
#     # Create inline keyboard
#     keyboard = InlineKeyboardMarkup()
#     keyboard.add(InlineKeyboardButton("Personal Info", callback_data='personal_info'))
#     keyboard.add(InlineKeyboardButton("Education", callback_data='education'))
#     keyboard.add(InlineKeyboardButton("Skills", callback_data='skills'))
#     keyboard.add(InlineKeyboardButton("Languages", callback_data='languages'))
#     keyboard.add(InlineKeyboardButton("Experience", callback_data='experience'))
#
#     bot.send_message(chat_id, "Choose a section:", reply_markup=keyboard)
#
#
# @bot.callback_query_handler(func=lambda call: True)
# def callback_query(call):
#     chat_id = call.message.chat.id
#     if call.data == 'personal_info':
#         send_personal_info(chat_id)
#     elif call.data == 'education':
#         send_education_info(chat_id)
#     elif call.data == 'skills':
#         send_skills_info(chat_id)
#     elif call.data == 'languages':
#         send_languages_info(chat_id)
#     elif call.data == 'experience':
#         send_experience_info(chat_id)
#
#
# def send_personal_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id,
#                          f"Personal Info:\nName: {user.name} {user.surname}\nAge: {user.age}\nAddress: {user.address}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
#
# def send_education_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id, f"Education Info:\n{user.education}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
#
# def send_skills_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id, f"Skills Info:\n{', '.join(user.skills.split(','))}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
#
# def send_languages_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id, f"Languages Info:\n{', '.join(user.languages.split(','))}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
#
# def send_experience_info(chat_id):
#     if chat_id in user_data:
#         user = user_data[chat_id]
#         bot.send_message(chat_id, f"Experience:\n{user.experience}")
#     else:
#         bot.send_message(chat_id, "User data not found. Please complete the initial steps.")
#
# def create_cv(message):
#     try:
#         cv_doc = Document()
#     except FileNotFoundError:
#         bot.send_message(message.chat.id, "The base document 'Resume.docx' was not found.")
#         return
#
#     chat_id = message.chat.id
#
#     # Add photo to the document
#     if 'Photo' in user_data[chat_id]:
#         photo_emoji = user_data[chat_id]['Photo']
#         cv_doc.add_paragraph(f"Photo: {photo_emoji}")
#
#     # Add header to the document
#     cv_doc.add_heading('Curriculum Vitae', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#     # Add user data to the document
#     for key, value in user_data[chat_id].items():
#         if key != 'Photo':
#             cv_doc.add_paragraph(f"{key}: {value}")
#
#     # Generate a clean file name without special characters
#     clean_name = "".join([c for c in user_data[chat_id]['Name'] if c.isalnum()])
#     clean_surname = "".join([c for c in user_data[chat_id]['Surname'] if c.isalnum()])
#     updated_file_name = f"{clean_name}_{clean_surname}_Updated_CV.docx"
#
#     cv_doc.save(updated_file_name)
#
#     with open(updated_file_name, 'rb') as doc:
#         bot.send_document(chat_id, doc)
#
#     user_data.pop(chat_id, None)
#
# bot.polling(none_stop=True)
